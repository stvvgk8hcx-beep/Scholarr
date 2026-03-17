"""Microsoft Word / Office document integration.

Extracts metadata and assignment information from .docx files.
Uses python-docx library for reading Word document properties and content.

Can extract:
  - Document metadata (title, author, subject, keywords, created/modified dates)
  - Headings and structure
  - Tables
  - Attempts to detect assignment info (course code, assignment number, etc.)
"""

from datetime import datetime
from typing import Any, Dict, List, Optional
import logging
import re
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from scholarr.core.integrations import BaseIntegrationProvider, IntegrationType, IntegrationStatus

logger = logging.getLogger(__name__)


class WordIntegrationProvider(BaseIntegrationProvider):
    """Word document integration provider.

    Extracts metadata and content from .docx files.
    Actual implementation (unlike LMS providers) since python-docx is available.
    """

    def __init__(self):
        """Initialize Word provider."""
        super().__init__("word", IntegrationType.DOCUMENT_TOOL)
        if not HAS_DOCX:
            logger.warning("python-docx not installed; Word integration will be limited")

    async def connect(self, config: Dict[str, Any]) -> bool:
        """Word documents don't need a persistent connection.

        Config can be empty or contain default document locations.

        Args:
            config: Configuration (can be empty)

        Returns:
            True (always succeeds)
        """
        self._is_connected = True
        self._clear_last_error()
        return True

    async def disconnect(self) -> bool:
        """Disconnect from Word provider.

        Returns:
            True (always succeeds)
        """
        self._is_connected = False
        return True

    async def sync(self) -> Dict[str, Any]:
        """Not applicable for document tool.

        Returns:
            Empty results
        """
        return {"message": "Word provider does not support sync"}

    async def test_connection(self) -> bool:
        """Word provider is always available if python-docx is installed.

        Returns:
            True if python-docx is available
        """
        return HAS_DOCX

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from a Word document.

        Args:
            file_path: Path to .docx file

        Returns:
            Dictionary with title, author, subject, keywords, created_date, modified_date
        """
        if not HAS_DOCX:
            return {"error": "python-docx not installed"}

        try:
            doc = Document(file_path)
            props = doc.core_properties

            return {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "created": props.created.isoformat() if props.created else None,
                "modified": props.modified.isoformat() if props.modified else None,
                "comments": props.comments or "",
                "file_path": file_path,
                "extracted_at": datetime.now().isoformat()
            }
        except Exception as e:
            logger.error(f"Failed to extract metadata from {file_path}: {e}")
            return {"error": str(e), "file_path": file_path}

    def extract_headings(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract heading structure from a Word document.

        Useful for parsing assignment instructions organized by sections.

        Args:
            file_path: Path to .docx file

        Returns:
            List of dicts with heading text and level
        """
        if not HAS_DOCX:
            return []

        try:
            doc = Document(file_path)
            headings = []

            for para in doc.paragraphs:
                if para.style.name.startswith("Heading"):
                    # Extract level from style name (Heading 1, Heading 2, etc.)
                    level = 1
                    if "Heading" in para.style.name:
                        parts = para.style.name.split()
                        if len(parts) > 1 and parts[1].isdigit():
                            level = int(parts[1])

                    headings.append({
                        "text": para.text.strip(),
                        "level": level,
                        "style": para.style.name
                    })

            return headings

        except Exception as e:
            logger.error(f"Failed to extract headings from {file_path}: {e}")
            return []

    def extract_tables(self, file_path: str) -> List[List[List[str]]]:
        """Extract all tables from a Word document.

        Tables often contain grade rubrics, requirements, etc.

        Args:
            file_path: Path to .docx file

        Returns:
            List of tables, each table is a list of rows, each row is a list of cell texts
        """
        if not HAS_DOCX:
            return []

        try:
            doc = Document(file_path)
            tables = []

            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            return tables

        except Exception as e:
            logger.error(f"Failed to extract tables from {file_path}: {e}")
            return []

    def extract_full_text(self, file_path: str) -> str:
        """Extract all text from a Word document.

        Args:
            file_path: Path to .docx file

        Returns:
            Full document text joined with newlines
        """
        if not HAS_DOCX:
            return ""

        try:
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            return ""

    def detect_assignment_info(self, file_path: str) -> Dict[str, Any]:
        """Try to detect assignment information from document content.

        Looks for common patterns:
          - Course code (e.g., COMP101, MATH201)
          - Assignment number (e.g., Assignment 1, Assign #3)
          - Due date
          - Submission type

        Args:
            file_path: Path to .docx file

        Returns:
            Dictionary with detected info (may have empty/None fields)
        """
        if not HAS_DOCX:
            return {}

        try:
            detected = {
                "course_code": None,
                "assignment_number": None,
                "title": None,
                "due_date": None,
                "submission_type": None,
                "confidence": "low"
            }

            # Get full text
            full_text = self.extract_full_text(file_path)
            metadata = self.extract_metadata(file_path)

            # Try to get title from metadata or document
            detected["title"] = metadata.get("title") or ""

            # Look for course code patterns (e.g., COMP101, CS-201, MATH 301)
            course_match = re.search(r'\b([A-Z]{2,})\s*(-?)(\d{3,4})\b', full_text)
            if course_match:
                detected["course_code"] = f"{course_match.group(1)}{course_match.group(3)}"

            # Look for assignment number/name patterns
            assign_match = re.search(
                r'(?:assignment|assign|homework|hw|problem set|ps)\s*(?:#|no\.?)?\s*(\d+)',
                full_text,
                re.IGNORECASE
            )
            if assign_match:
                detected["assignment_number"] = int(assign_match.group(1))

            # Look for due date patterns
            due_match = re.search(
                r'due\s*(?:date|on)?:?\s*(\d{1,2}/\d{1,2}/\d{2,4}|\w+\s+\d{1,2}(?:st|nd|rd|th)?(?:,?\s*\d{4})?)',
                full_text,
                re.IGNORECASE
            )
            if due_match:
                detected["due_date"] = due_match.group(1)

            # Look for submission type
            if re.search(r'(?:upload|submit|attach).*?(?:file|document|pdf|word|doc)', full_text, re.IGNORECASE):
                detected["submission_type"] = "file_upload"
            elif re.search(r'(?:submit|enter).*?(?:text|answer|response|essay)', full_text, re.IGNORECASE):
                detected["submission_type"] = "text_entry"
            elif re.search(r'(?:link|url|submit).*?(?:link|url)', full_text, re.IGNORECASE):
                detected["submission_type"] = "url_submission"

            # Simple heuristic for confidence
            info_found = sum([
                detected["course_code"] is not None,
                detected["assignment_number"] is not None,
                detected["due_date"] is not None,
                detected["submission_type"] is not None
            ])
            if info_found >= 3:
                detected["confidence"] = "high"
            elif info_found >= 2:
                detected["confidence"] = "medium"

            return detected

        except Exception as e:
            logger.error(f"Failed to detect assignment info from {file_path}: {e}")
            return {"error": str(e)}

    async def get_status(self) -> IntegrationStatus:
        """Get status of Word provider.

        Returns:
            IntegrationStatus
        """
        return IntegrationStatus(
            provider_name="word",
            provider_type=IntegrationType.DOCUMENT_TOOL,
            is_connected=self._is_connected,
            last_sync=self._last_sync,
            last_error=self._last_error,
            configuration_valid=HAS_DOCX,
            metadata={
                "python_docx_available": HAS_DOCX,
                "capabilities": ["extract_metadata", "extract_headings", "extract_tables", "detect_assignment_info"]
            }
        )
